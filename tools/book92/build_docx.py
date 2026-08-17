# -*- coding: utf-8 -*-
import re, os, sys, json
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.chdir('/home/claude/b90')
PAGEMAP = {}
if len(sys.argv) > 1 and sys.argv[1] == 'withpage' and os.path.exists('pagemap.json'):
    PAGEMAP = json.load(open('pagemap.json', encoding='utf-8'))

SERIF = 'Noto Serif CJK SC'
SANS = 'Noto Sans CJK SC'
DARK = RGBColor(0x1D, 0x39, 0x55)
GREY = RGBColor(0x6B, 0x7B, 0x8C)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(170), Mm(240)
sec.top_margin, sec.bottom_margin = Mm(17), Mm(15)
sec.left_margin, sec.right_margin = Mm(16), Mm(15)
sec.footer_distance = Mm(8)
sec.different_first_page_header_footer = True

st = doc.styles['Normal']
st.font.name = SERIF
st.font.size = Pt(9.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
pf = st.paragraph_format
pf.line_spacing = 1.35
pf.space_after = Pt(1.5)
pf.first_line_indent = Pt(19)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _fld(par, instr):
    r = par.add_run()
    a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
    b = OxmlElement('w:instrText'); b.set(qn('xml:space'), 'preserve'); b.text = instr
    c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'end')
    r._r.append(a); r._r.append(b); r._r.append(c)


def add_footer():
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    _fld(p, 'PAGE')
    for r in p.runs:
        r.font.size = Pt(8.5); r.font.name = SANS
        r._element.rPr.rFonts.set(qn('w:eastAsia'), SANS)
        r.font.color.rgb = GREY


def add_header():
    p = sec.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run('谁 来 陪 伴 我 ？')
    r.font.size = Pt(8); r.font.name = SANS
    r._element.rPr.rFonts.set(qn('w:eastAsia'), SANS)
    r.font.color.rgb = GREY
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '4')
    bt.set(qn('w:space'), '1'); bt.set(qn('w:color'), 'B7C3CE')
    pbdr.append(bt); pPr.append(pbdr)


add_footer(); add_header()


def P(text='', size=9.5, bold=False, align=None, indent=None, font=SERIF,
      color=None, before=0, after=1.5, pagebreak=False, spacing=1.35):
    p = doc.add_paragraph()
    if pagebreak:
        p.paragraph_format.page_break_before = True
    f = p.paragraph_format
    f.space_before = Pt(before); f.space_after = Pt(after)
    f.line_spacing = spacing
    f.first_line_indent = Pt(indent if indent is not None else 19)
    if align is not None:
        f.alignment = align
    if text:
        for seg, bd in parse_bold(text):
            r = p.add_run(seg)
            r.font.size = Pt(size); r.font.name = font
            r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            r.bold = bold or bd
            if color is not None:
                r.font.color.rgb = color
    return p


def parse_bold(t):
    out = []
    for i, seg in enumerate(re.split(r'\*\*', t)):
        if seg:
            out.append((seg, i % 2 == 1))
    return out or [('', False)]


def clean(t):
    t = t.replace('\u2003', '　')
    return t


def toc_entry(par, title, level):
    if not PAGEMAP:
        return
    pg = PAGEMAP.get(title)
    if not pg:
        return
    par.paragraph_format.tab_stops.add_tab_stop(Inches(5.42), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = par.add_run('\t' + str(pg))
    r.font.size = Pt(9); r.font.name = SANS
    r._element.rPr.rFonts.set(qn('w:eastAsia'), SANS)
    r.font.color.rgb = GREY


src = open('manuscript.md', encoding='utf-8').read()
lines = src.split('\n')

# ---------- 扉页 ----------
P(); P(); P(); P()
P('谁来陪伴我？', size=30, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, color=DARK, after=10)
P('AI 时代的婚姻困境', size=13, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, color=GREY, after=40)
P('王德生　＋　Claude　编著', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, after=6)
P('德麦国际出版社', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, color=GREY, after=2)
P('专著第 92 号', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, color=GREY)

# ---------- 目录占位（第二遍才有页码） ----------
TOC_TITLES = []
for l in lines:
    if l.startswith('## '):
        TOC_TITLES.append(('c', clean(l[3:].strip())))
    elif l.startswith('# ') and l[2:].startswith('第'):
        TOC_TITLES.append(('b', clean(l[2:].strip())))

P('目　录', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0,
  color=DARK, pagebreak=True, after=12)
for kind, t in TOC_TITLES:
    if t == '封　底':
        continue
    p = P('', indent=0, after=2.5)
    r = p.add_run(('　　' if kind == 'c' else '') + t)
    r.font.size = Pt(9.5 if kind == 'c' else 10)
    r.font.name = SERIF
    r._element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
    r.bold = (kind == 'b')
    if kind == 'b':
        r.font.color.rgb = DARK
    toc_entry(p, t, kind)

# ---------- 正文 ----------
i = 0
in_table = []


def flush_table():
    global in_table
    if not in_table:
        return
    rows = [r for r in in_table if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', r)]
    cells = [[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
    ncol = max(len(c) for c in cells)
    tb = doc.add_table(rows=len(cells), cols=ncol)
    tb.style = 'Table Grid'
    for ri, row in enumerate(cells):
        for ci in range(ncol):
            cell = tb.cell(ri, ci)
            cell.text = ''
            par = cell.paragraphs[0]
            par.paragraph_format.first_line_indent = Pt(0)
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.line_spacing = 1.15
            txt = row[ci] if ci < len(row) else ''
            for seg, bd in parse_bold(txt):
                r = par.add_run(seg)
                r.font.size = Pt(8)
                r.font.name = SANS if ri == 0 else SERIF
                r._element.rPr.rFonts.set(qn('w:eastAsia'), SANS if ri == 0 else SERIF)
                r.bold = bd or (ri == 0)
                if ri == 0:
                    r.font.color.rgb = DARK
    in_table = []


while i < len(lines):
    l = clean(lines[i].rstrip())
    i += 1
    if l.startswith('|'):
        in_table.append(l); continue
    flush_table()
    if not l.strip():
        continue
    if l.startswith('# 谁来陪伴我'):
        # 跳过稿本自带书名页（前面已排）
        while i < len(lines) and not lines[i].startswith('## '):
            i += 1
        continue
    if l.startswith('# 第') and '编' in l[:6]:
        P('', pagebreak=True)
        P(); P(); P(); P(); P(); P()
        head = l[2:]
        bn, _, bt = head.partition('　')
        P(bn, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, color=GREY, after=10)
        P(bt, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, color=DARK, after=14)
        P('· · ·', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0,
          color=RGBColor(0xC9, 0xA2, 0x27))
        continue
    if l.startswith('## '):
        t = l[3:].strip()
        p = P('', pagebreak=True, after=0)
        m = re.match(r'^(第[一二三四五六七八九十]+章|枢纽章|合章)　(.*)$', t)
        if m:
            r = p.add_run(m.group(1))
            r.font.size = Pt(10); r.font.name = SANS
            r._element.rPr.rFonts.set(qn('w:eastAsia'), SANS)
            r.font.color.rgb = GREY
            p.paragraph_format.first_line_indent = Pt(0)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
            bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '4')
            bt.set(qn('w:space'), '4'); bt.set(qn('w:color'), 'C9CFD6')
            pbdr.append(bt); pPr.append(pbdr)
            P(m.group(2), size=16.5, bold=True, indent=0, color=DARK, before=8, after=12)
        else:
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(t)
            r.font.size = Pt(16.5); r.font.name = SERIF
            r._element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
            r.bold = True; r.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(12)
        continue
    if l.startswith('#### '):
        P('▍' + l[5:].strip(), size=10, bold=True, indent=0, before=9, after=3, color=DARK)
        continue
    if l.startswith('### '):
        P('▍' + l[4:].strip(), size=11, bold=True, indent=0, before=11, after=4, color=DARK)
        continue
    if l.startswith('> '):
        p = P(l[2:].strip(), size=9.5, indent=0, before=5, after=5)
        p.paragraph_format.left_indent = Pt(22)
        p.paragraph_format.right_indent = Pt(10)
        for r in p.runs:
            r.font.color.rgb = DARK
        continue
    if l.startswith('- ') or l.startswith('* '):
        p = P('· ' + l[2:].strip(), size=9, indent=0, after=1)
        p.paragraph_format.left_indent = Pt(14)
        continue
    if l.startswith('---'):
        continue
    P(l)

flush_table()
doc.save('谁来陪伴我.docx')
print('docx saved')
