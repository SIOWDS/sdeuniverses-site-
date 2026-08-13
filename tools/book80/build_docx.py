# -*- coding: utf-8 -*-
"""《判断的危机》docx 构建 · 170×240mm · 完整体例"""
import re, os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = '/home/claude/book80/判断的危机.md'
OUT = '/home/claude/book80/判断的危机.docx'
SERIF, SANS = 'Noto Serif CJK SC', 'Noto Sans CJK SC'
INK, ACC, ACC2 = RGBColor(0x1A, 0x1A, 0x1A), RGBColor(0x1D, 0x39, 0x55), RGBColor(0x7A, 0x93, 0xAE)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(170), Mm(240)
sec.top_margin, sec.bottom_margin = Mm(17), Mm(15)
sec.left_margin, sec.right_margin = Mm(16), Mm(15)
sec.footer_distance = Mm(8)

st = doc.styles['Normal']
st.font.name, st.font.size, st.font.color.rgb = SERIF, Pt(9.5), INK
st.element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
pf = st.paragraph_format
pf.line_spacing, pf.space_after, pf.first_line_indent = 1.35, Pt(0), Pt(19)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

PENDING = {'brk': False}


def setfont(run, name=SERIF, size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size: run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None: run.font.color.rgb = color


def para(text='', size=None, bold=False, align=None, indent=None,
         before=0, after=1.5, color=None, name=SERIF, spacing=None):
    p = doc.add_paragraph()
    if PENDING['brk']:
        p.paragraph_format.page_break_before = True
        PENDING['brk'] = False
    f = p.paragraph_format
    f.space_before, f.space_after = Pt(before), Pt(after)
    f.first_line_indent = Pt(19 if indent is None else indent)
    if align is not None: f.alignment = align
    if spacing: f.line_spacing = spacing
    if text:
        setfont(p.add_run(text), name, size, bold, color)
    return p


def pagebreak(): PENDING['brk'] = True


GOLD = RGBColor(0xB8, 0x8A, 0x3E)


def rule(color='C7D2DC', sz='6'):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), color)
    pbdr.append(bot); pPr.append(pbdr)


def set_cell_border(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'left', 'right'):
        spec = kw.get(edge)
        el = OxmlElement('w:' + edge)
        if spec is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec[0]))
            el.set(qn('w:color'), spec[1])
        borders.append(el)
    tcPr.append(borders)


def add_header():
    sec.different_first_page_header_footer = True
    p = sec.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    setfont(p.add_run('判 断 的 危 机'), SANS, 8, False, ACC2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), 'C7D2DC')
    pbdr.append(bot); pPr.append(pbdr)


def add_footer():
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(); setfont(r, SANS, 8.5, False, ACC2)
    for el, attr in (('w:fldChar', {'w:fldCharType': 'begin'}), ('w:instrText', None),
                     ('w:fldChar', {'w:fldCharType': 'end'})):
        e = OxmlElement(el)
        if attr:
            for k, v in attr.items(): e.set(qn(k), v)
        else:
            e.set(qn('xml:space'), 'preserve'); e.text = ' PAGE '
        r._element.append(e)


def table(rows):
    """rows: list of list[str]，第一行为表头"""
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    n = len(rows)
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            set_cell_border(c,
                            top=(12, '1D3955') if i == 0 else None,
                            bottom=(6, '7A93AE') if i == 0 else ((12, '1D3955') if i == n - 1 else None))
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.2
            setfont(p.add_run(cell), SANS if i == 0 else SERIF, 8.5, i == 0, ACC if i == 0 else INK)
    para('', after=4)


raw = open(SRC, encoding='utf-8').read()
raw = re.sub(r'\*\*(.+?)\*\*', r'\1', raw).replace('**', '')
blocks = [b.strip() for b in raw.split('\n\n') if b.strip()]

# ---------- 扉页 ----------
para('', after=105)
para('判 断 的 危 机', size=34, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0,
     color=ACC, spacing=1.2)
para('它不会发出任何警报', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, before=14, color=ACC2)
para('', after=88)
para('王德生　＋　Claude', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, before=58)
para('德麦国际出版社', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, before=10, color=ACC2)
para('SDE UNIVERSES', size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, before=4,
     color=ACC2, name=SANS)
pagebreak()

# ---------- 出版信息页 ----------
para('出版信息', size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, after=16, color=ACC)
INFO = [('书　　名', '判断的危机——它不会发出任何警报'),
        ('著　　者', '王德生 ＋ Claude'),
        ('出版发行', '德麦国际出版社　Demai International Press'),
        ('版　　次', '2026 年 8 月第 1 版第 1 次印刷'),
        ('开　　本', '16 开　170mm × 240mm'),
        ('字　　数', '约 23.6 万汉字'),
        ('专著编号', '德麦国际专著第 80 号'),
        ('I S B N', '申领中'),
        ('分　　类', '人工智能 / 制度经济 / 科学社会学 / 教育评价')]
for k, v in INFO:
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    setfont(p.add_run(k + '　'), SANS, 9.5, True, ACC)
    setfont(p.add_run(v), SERIF, 9.5)
para('', after=10)
for s in ['本书由「学科通融」专栏之四十四至之五十三共十篇连续研究编成，另新写前言、导读、导论、枢纽章、合章、结语与附录。十篇原刊于 sdeuniverses.com，作者王德生＋Claude；编入本书时篇名改为章名、篇际互引改为章际互引，正文未作删节。',
          '本书的前言、导读、导论、枢纽章的书内导入与推广、合章、结语、附录与后记由编者写作，因此按本栏规程，编者不得为本书出具创新智商认证分；全书盲评待未参与写作的一方独立完成后公布。',
          '版权所有　侵权必究。本书内容可自由引用与批评，引用时请注明出处。']:
    para(s, after=7)
pagebreak()


def is_h1(b): return b.startswith('# ')
def is_h3(b): return b.startswith('### ')


# ---------- 目录 ----------
para('目　录', size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, after=16, color=ACC)
for b in blocks:
    if is_h1(b):
        t = b[2:].strip()
        isb = t.startswith('编')
        ischap = t.startswith('第') and '章' in t
        para(t, size=11 if isb else 10, bold=isb, indent=14 if ischap else 0,
             after=4, before=7 if isb else 0, color=ACC if isb else INK)
pagebreak()
add_footer()
add_header()

# ---------- 正文 ----------
first = False
inback = False
i = 0
while i < len(blocks):
    b = blocks[i]
    if is_h1(b):
        t = b[2:].strip()
        if not first: pagebreak()
        first = False
        inback = t.startswith('封')
        if t.startswith('编'):
            para('', after=112)
            head, _, tail = t.partition('·')
            para(head.strip(), size=13, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0,
                 color=ACC2, name=SANS, after=10)
            para(tail.strip(), size=25, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0,
                 color=ACC, spacing=1.3)
            para('· · ·', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0,
                 color=GOLD, before=16, name=SANS)
        else:
            head, _, tail = t.partition('·')
            para('', after=16)
            if tail.strip():
                para(head.strip(), size=10, indent=0, color=ACC2, name=SANS, after=5)
                rule()
                para(tail.strip(), size=16.5, bold=True, indent=0, after=13, before=7,
                     color=ACC, spacing=1.25)
            else:
                para(head.strip(), size=16.5, bold=True, indent=0, after=13, before=7,
                     color=ACC, spacing=1.25)
    elif is_h3(b):
        para('▍' + b[4:].strip(), size=11, bold=True, indent=0, before=11, after=4, color=ACC)
    elif b.startswith('|') and '\n' in b:
        rows = []
        for line in b.split('\n'):
            line = line.strip()
            if not line.startswith('|'): continue
            cells = [c.strip() for c in line.strip('|').split('|')]
            if all(set(c) <= set('-: ') for c in cells): continue
            rows.append(cells)
        if rows and len({len(r) for r in rows}) == 1:
            table(rows)
        else:
            for r in rows: para('　'.join(r), after=1.5)
    elif b.startswith('- '):
        for line in b.split('\n'):
            para('· ' + line[2:].strip(), indent=0, after=1)
    elif b.startswith('---'):
        pass
    elif inback:
        para(b.replace('\n', ''), after=8, indent=0, align=WD_ALIGN_PARAGRAPH.CENTER,
             size=10.5, color=ACC if b.startswith('判断的危机') else None,
             bold=b.startswith('判断的危机'), spacing=1.5)
    else:
        para(b.replace('\n', ''), after=1.5)
    i += 1

doc.save(OUT)
print('saved', OUT, os.path.getsize(OUT))
